from __future__ import annotations

import argparse
import json
import re
import time
import urllib.parse
import urllib.request
from dataclasses import asdict, dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any


DEFAULT_ANTI_API = "http://127.0.0.1:8000"
DEFAULT_BAIBAI_API = "http://127.0.0.1:18765"
DEFAULT_OUTPUT_ROOT = Path(r"<PRIVATE_THESIS_WORKSPACE>\pipeline_outputs")
DEFAULT_SEGMENT_MAX_CHARS = 800
DEFAULT_THRESHOLD = 0.4
DEFAULT_MAX_REVIEW_PASSES = 6


@dataclass
class ReviewStep:
    pass_index: int
    action: str
    input_chars: int
    risk: float
    reasons: list[str]
    baibai_round_outputs: list[str] = field(default_factory=list)


@dataclass
class SegmentRecord:
    index: int
    segment_id: str
    text: str
    initial_risk: float | None
    initial_reasons: list[str]
    selected: bool
    processed_text: str
    final_risk: float | None = None
    final_reasons: list[str] | None = None
    baibai_round_outputs: list[str] | None = None
    pass_count: int = 0
    unresolved: bool = False
    unresolved_reason: str | None = None
    processing_error: str | None = None
    review_steps: list[ReviewStep] = field(default_factory=list)


def post_json(url: str, payload: dict[str, Any], timeout: int = 180) -> dict[str, Any]:
    request = urllib.request.Request(
        url,
        data=json.dumps(payload, ensure_ascii=False).encode("utf-8"),
        headers={"Content-Type": "application/json"},
        method="POST",
    )
    with urllib.request.urlopen(request, timeout=timeout) as response:
        return json.loads(response.read().decode("utf-8"))


def get_json(url: str, timeout: int = 60) -> dict[str, Any]:
    with urllib.request.urlopen(url, timeout=timeout) as response:
        return json.loads(response.read().decode("utf-8"))


def read_document(path: Path, max_chars: int = DEFAULT_SEGMENT_MAX_CHARS) -> tuple[str, list[str]]:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        text = path.read_text(encoding="utf-8")
        return text, split_text_to_segments(text, max_chars=max_chars)
    if suffix == ".docx":
        try:
            from docx import Document  # type: ignore[import]
        except ImportError as exc:
            raise RuntimeError("Reading .docx requires python-docx. Run with baibaiAIGC\\.venv\\Scripts\\python.exe.") from exc
        document = Document(str(path))
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        return "\n\n".join(paragraphs), build_segment_groups(paragraphs, max_chars=max_chars)
    raise RuntimeError(f"Unsupported input file type: {path.suffix}")


def split_text_to_segments(text: str, max_chars: int = DEFAULT_SEGMENT_MAX_CHARS) -> list[str]:
    paragraphs = [part.strip() for part in re.split(r"\n\s*\n", text.replace("\r\n", "\n")) if part.strip()]
    return build_segment_groups(paragraphs, max_chars=max_chars)


def split_long_paragraph(paragraph: str, max_chars: int = DEFAULT_SEGMENT_MAX_CHARS) -> list[str]:
    if len(paragraph) <= max_chars:
        return [paragraph]

    sentence_parts = [part.strip() for part in re.split(r"(?<=[。！？；;.!?])", paragraph) if part.strip()]
    if len(sentence_parts) <= 1:
        return hard_wrap_text(paragraph, max_chars=max_chars)

    chunks: list[str] = []
    current: list[str] = []
    current_len = 0
    for sentence in sentence_parts:
        if len(sentence) > max_chars:
            if current:
                chunks.append("".join(current).strip())
                current = []
                current_len = 0
            chunks.extend(hard_wrap_text(sentence, max_chars=max_chars))
            continue
        if current and current_len + len(sentence) > max_chars:
            chunks.append("".join(current).strip())
            current = [sentence]
            current_len = len(sentence)
        else:
            current.append(sentence)
            current_len += len(sentence)
    if current:
        chunks.append("".join(current).strip())
    return chunks


def hard_wrap_text(text: str, max_chars: int = DEFAULT_SEGMENT_MAX_CHARS) -> list[str]:
    return [text[start : start + max_chars].strip() for start in range(0, len(text), max_chars) if text[start : start + max_chars].strip()]


def build_segment_groups(paragraphs: list[str], max_chars: int = DEFAULT_SEGMENT_MAX_CHARS) -> list[str]:
    groups: list[str] = []
    current: list[str] = []
    current_len = 0
    for raw_paragraph in paragraphs:
        for paragraph in split_long_paragraph(raw_paragraph, max_chars=max_chars):
            join_cost = 2 if current else 0
            paragraph_len = len(paragraph)
            if current and current_len + join_cost + paragraph_len > max_chars:
                groups.append("\n\n".join(current))
                current = [paragraph]
                current_len = paragraph_len
            else:
                current.append(paragraph)
                current_len += join_cost + paragraph_len
    if current:
        groups.append("\n\n".join(current))
    return groups


def detect_with_anti(anti_api: str, text: str, threshold_profile: str = "balanced") -> dict[str, Any]:
    return post_json(
        f"{anti_api.rstrip('/')}/api/detect",
        {
            "rawText": text,
            "language": "zh-CN",
            "segmentMode": "auto",
            "thresholdProfile": threshold_profile,
        },
        timeout=90,
    )


def risk_summary(detection: dict[str, Any]) -> tuple[float, list[str]]:
    segment_results = detection.get("segmentResults")
    if not isinstance(segment_results, list) or not segment_results:
        return float(detection.get("overallRisk", 0.0)), []
    max_item = max(segment_results, key=lambda item: float(item.get("riskScore", 0.0)))
    reasons = max_item.get("reasons", [])
    return float(max_item.get("riskScore", detection.get("overallRisk", 0.0))), [str(item) for item in reasons]


def upload_baibai_text(baibai_api: str, filename: str, text: str) -> str:
    payload = post_json(
        f"{baibai_api.rstrip('/')}/api/upload-document",
        {"filename": filename, "encoding": "text", "content": text},
        timeout=90,
    )
    return str(payload["sourcePath"])


def should_retry_baibai_error(message: str) -> bool:
    transient_tokens = [
        "status 502",
        "status 503",
        "status 504",
        "status 500",
        "upstream error",
        "do_request_failed",
        "do request failed",
        "timed out",
        "timeout",
        "temporarily",
        "10054",
        "forcibly closed",
        "远程主机强迫关闭",
        "empty chat/completions message.content",
        "did not return usable text",
        "disallowed answer-style pattern",
        "answer-style drift",
    ]
    lower_message = message.lower()
    return any(token in lower_message for token in transient_tokens)


def run_baibai_round(
    baibai_api: str,
    source_path: str,
    model_config: dict[str, Any],
    timeout: int = 900,
    max_attempts: int = 3,
) -> dict[str, Any]:
    last_error: Exception | None = None
    for attempt in range(1, max_attempts + 1):
        try:
            return run_baibai_round_once(baibai_api, source_path, model_config, timeout=timeout)
        except Exception as exc:  # noqa: BLE001 - keep the outer pipeline resilient to gateway errors.
            last_error = exc
            if attempt >= max_attempts or not should_retry_baibai_error(str(exc)):
                raise
            time.sleep(min(5 * attempt, 20))
    if last_error:
        raise last_error
    raise RuntimeError("baibai round failed before starting.")


def run_baibai_round_once(baibai_api: str, source_path: str, model_config: dict[str, Any], timeout: int = 900) -> dict[str, Any]:
    start = post_json(
        f"{baibai_api.rstrip('/')}/api/run-round",
        {"sourcePath": source_path, "modelConfig": model_config},
        timeout=90,
    )
    run_id = str(start["runId"])
    event_url = f"{baibai_api.rstrip('/')}/api/run-round-events/{run_id}"
    request = urllib.request.Request(event_url, method="GET")
    deadline = time.time() + timeout
    with urllib.request.urlopen(request, timeout=timeout) as response:
        current_event: str | None = None
        for raw_line in response:
            if time.time() > deadline:
                raise TimeoutError(f"baibai round timed out: {run_id}")
            line = raw_line.decode("utf-8", errors="replace").strip()
            if not line:
                continue
            if line.startswith("event:"):
                current_event = line.split(":", 1)[1].strip()
                continue
            if line.startswith("data:"):
                payload = json.loads(line.split(":", 1)[1].strip())
                if current_event == "error":
                    raise RuntimeError(f"baibai round failed: {payload}")
                if current_event == "result":
                    return payload
    raise RuntimeError(f"baibai round produced no result event: {run_id}")


def read_baibai_output(baibai_api: str, output_path: str) -> str:
    query = urllib.parse.quote(output_path)
    payload = get_json(f"{baibai_api.rstrip('/')}/api/read-output?outputPath={query}", timeout=90)
    return str(payload["text"])


def process_with_baibai(
    baibai_api: str,
    text: str,
    filename: str,
    rounds: int,
    model_config: dict[str, Any] | None = None,
) -> tuple[str, list[str]]:
    resolved_model_config = model_config or get_json(f"{baibai_api.rstrip('/')}/api/model-config", timeout=60)
    source_path = upload_baibai_text(baibai_api, filename, text)
    round_outputs: list[str] = []
    result: dict[str, Any] | None = None
    for _ in range(rounds):
        result = run_baibai_round(baibai_api, source_path, resolved_model_config)
        round_outputs.append(str(result["outputPath"]))
    if not result:
        return text, []
    return read_baibai_output(baibai_api, str(result["outputPath"])), round_outputs


def process_segment_with_review(
    *,
    index: int,
    segment: str,
    run_id: str,
    anti_api: str,
    baibai_api: str,
    model_config: dict[str, Any],
    rounds: int,
    threshold: float,
    max_review_passes: int,
    retry_max_chars: int,
) -> SegmentRecord:
    initial_detection = detect_with_anti(anti_api, segment)
    initial_risk, initial_reasons = risk_summary(initial_detection)
    review_steps = [
        ReviewStep(
            pass_index=0,
            action="initial_anti_only",
            input_chars=len(segment),
            risk=initial_risk,
            reasons=initial_reasons,
            baibai_round_outputs=[],
        )
    ]

    if initial_risk < threshold:
        return SegmentRecord(
            index=index,
            segment_id=f"seg-{index}",
            text=segment,
            initial_risk=initial_risk,
            initial_reasons=initial_reasons,
            selected=False,
            processed_text=segment,
            final_risk=initial_risk,
            final_reasons=initial_reasons,
            baibai_round_outputs=[],
            pass_count=0,
            unresolved=False,
            unresolved_reason=None,
            processing_error=None,
            review_steps=review_steps,
        )

    current_text = segment
    current_risk = initial_risk
    current_reasons = initial_reasons
    all_round_outputs: list[str] = []
    pass_count = 0
    unresolved = False
    unresolved_reason: str | None = None

    while current_risk >= threshold:
        if pass_count >= max_review_passes:
            unresolved = True
            unresolved_reason = f"risk remained >= {threshold} after {max_review_passes} baibai pass(es)"
            break
        if len(current_text) > retry_max_chars:
            unresolved = True
            unresolved_reason = f"segment is longer than retry max chars ({len(current_text)} > {retry_max_chars})"
            break

        next_pass = pass_count + 1
        try:
            processed_text, round_outputs = process_with_baibai(
                baibai_api,
                current_text,
                f"{run_id}_seg{index:03d}_pass{next_pass:02d}.txt",
                rounds,
                model_config=model_config,
            )
        except Exception as exc:  # noqa: BLE001 - keep the chapter run resumable when one segment fails.
            unresolved = True
            unresolved_reason = f"baibai pass {next_pass} failed"
            return SegmentRecord(
                index=index,
                segment_id=f"seg-{index}",
                text=segment,
                initial_risk=initial_risk,
                initial_reasons=initial_reasons,
                selected=True,
                processed_text=current_text,
                final_risk=current_risk,
                final_reasons=current_reasons,
                baibai_round_outputs=all_round_outputs,
                pass_count=pass_count,
                unresolved=unresolved,
                unresolved_reason=unresolved_reason,
                processing_error=str(exc),
                review_steps=review_steps,
            )

        current_text = processed_text
        all_round_outputs.extend(round_outputs)
        pass_count = next_pass
        detection = detect_with_anti(anti_api, current_text)
        current_risk, current_reasons = risk_summary(detection)
        review_steps.append(
            ReviewStep(
                pass_index=pass_count,
                action="baibai_two_rounds_then_anti",
                input_chars=len(current_text),
                risk=current_risk,
                reasons=current_reasons,
                baibai_round_outputs=round_outputs,
            )
        )

    return SegmentRecord(
        index=index,
        segment_id=f"seg-{index}",
        text=segment,
        initial_risk=initial_risk,
        initial_reasons=initial_reasons,
        selected=True,
        processed_text=current_text,
        final_risk=current_risk,
        final_reasons=current_reasons,
        baibai_round_outputs=all_round_outputs,
        pass_count=pass_count,
        unresolved=unresolved,
        unresolved_reason=unresolved_reason,
        review_steps=review_steps,
    )


def write_docx_if_possible(path: Path, paragraphs: list[str]) -> None:
    try:
        from docx import Document  # type: ignore[import]
    except ImportError:
        return
    document = Document()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    document.save(str(path))


def write_reports(
    output_dir: Path,
    input_path: Path,
    original_text: str,
    records: list[SegmentRecord],
    *,
    threshold: float,
    segment_max_chars: int,
    max_review_passes: int,
) -> None:
    output_dir.mkdir(parents=True, exist_ok=True)
    merged_segments = [record.processed_text for record in records]
    merged_text = "\n\n".join(merged_segments)

    (output_dir / "original.txt").write_text(original_text, encoding="utf-8")
    (output_dir / "merged_revised.txt").write_text(merged_text, encoding="utf-8")
    write_docx_if_possible(output_dir / "merged_revised.docx", merged_segments)

    report_payload = {
        "input_path": str(input_path),
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "workflow": "anti_first_then_iterative_baibai_review",
        "threshold": threshold,
        "segment_max_chars": segment_max_chars,
        "max_review_passes": max_review_passes,
        "segment_count": len(records),
        "anti_passed_without_baibai_count": sum(1 for record in records if record.pass_count == 0 and not record.unresolved),
        "initial_processed_count": sum(1 for record in records if record.pass_count > 0),
        "selected_count": sum(1 for record in records if record.selected),
        "reprocessed_count": sum(1 for record in records if record.pass_count > 1),
        "unresolved_count": sum(1 for record in records if record.unresolved),
        "segments": [asdict(record) for record in records],
    }
    (output_dir / "pipeline_report.json").write_text(
        json.dumps(report_payload, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    def format_risk(value: float | None) -> str:
        return "" if value is None else f"{value:.3f}"

    lines = [
        "# AntiAIGC -> baibaiAIGC -> AntiAIGC -> Iterative baibaiAIGC -> AWAS Review Packet",
        "",
        f"- Input: `{input_path}`",
        f"- Generated at: {report_payload['generated_at']}",
        f"- Segment max chars: {segment_max_chars}",
        f"- Risk threshold: {threshold}",
        f"- Max baibai review passes per segment: {max_review_passes}",
        f"- Segments: {report_payload['segment_count']}",
        f"- Passed initial AntiAIGC without baibai: {report_payload['anti_passed_without_baibai_count']}",
        f"- Entered baibai after initial AntiAIGC: {report_payload['initial_processed_count']}",
        f"- Needed 2+ baibai passes after review: {report_payload['reprocessed_count']}",
        f"- Still unresolved after safety limit / length gate: {report_payload['unresolved_count']}",
        "",
        "## Segment Summary",
        "",
        "| # | Baibai Passes | Entered baibai | Initial Risk | Final Risk | Unresolved | Main Reasons |",
        "|---|---:|---|---:|---:|---|---|",
    ]
    for record in records:
        initial_reasons = "; ".join(record.initial_reasons)
        final_risk = format_risk(record.final_risk)
        unresolved = record.unresolved_reason or record.processing_error or ""
        lines.append(
            f"| {record.index} | {record.pass_count} | {record.selected} | "
            f"{format_risk(record.initial_risk)} | {final_risk} | {unresolved} | {initial_reasons} |"
        )
    lines.extend(
        [
            "",
            "## AWAS Final Review Checklist",
            "",
            "- 中国药科大学专业硕士论文适配度",
            "- 章节逻辑与研究问题一致性",
            "- 方法、结果、讨论的证据-论点对应关系",
            "- 术语、缩略词、数据、引用是否保持稳定",
            "- 表题、图题、统计表达和 GB/T 7714 引用风险",
            "- 语言清晰度、学术语体、模板化表达残留",
            "- 人工核对优先级：先检查被处理片段，再检查前后段衔接",
        ]
    )
    (output_dir / "awas_review_packet.md").write_text("\n".join(lines), encoding="utf-8")


def parse_args(argv: list[str] | None = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Run baibaiAIGC -> AntiAIGC iterative review pipeline for a manuscript.")
    parser.add_argument("input", type=Path, help="Input .txt, .md, or .docx manuscript path.")
    parser.add_argument("--anti-api", default=DEFAULT_ANTI_API)
    parser.add_argument("--baibai-api", default=DEFAULT_BAIBAI_API)
    parser.add_argument("--output-root", type=Path, default=DEFAULT_OUTPUT_ROOT)
    parser.add_argument("--threshold", type=float, default=DEFAULT_THRESHOLD)
    parser.add_argument("--rounds", type=int, default=2, choices=[1, 2])
    parser.add_argument("--segment-max-chars", type=int, default=DEFAULT_SEGMENT_MAX_CHARS)
    parser.add_argument("--retry-max-chars", type=int, default=DEFAULT_SEGMENT_MAX_CHARS)
    parser.add_argument("--max-review-passes", type=int, default=DEFAULT_MAX_REVIEW_PASSES)
    parser.add_argument("--all", action="store_true", help="Deprecated; this workflow always processes every segment first.")
    parser.add_argument("--limit", type=int, default=0, help="Smoke-test only: process at most N initial segments; 0 means no limit.")
    return parser.parse_args(argv)


def main(argv: list[str] | None = None) -> int:
    args = parse_args(argv)
    input_path = args.input.resolve()
    if not input_path.exists():
        raise FileNotFoundError(input_path)
    if args.max_review_passes < 1:
        raise ValueError("--max-review-passes must be at least 1.")
    if args.segment_max_chars <= 0 or args.retry_max_chars <= 0:
        raise ValueError("--segment-max-chars and --retry-max-chars must be positive.")

    run_id = f"{input_path.stem}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    output_dir = args.output_root / run_id
    original_text, segments = read_document(input_path, max_chars=args.segment_max_chars)
    if args.limit > 0:
        segments = segments[: args.limit]
    model_config = get_json(f"{args.baibai_api.rstrip('/')}/api/model-config", timeout=60)

    records: list[SegmentRecord] = []
    for index, segment in enumerate(segments, start=1):
        record = process_segment_with_review(
            index=index,
            segment=segment,
            run_id=run_id,
            anti_api=args.anti_api,
            baibai_api=args.baibai_api,
            model_config=model_config,
            rounds=args.rounds,
            threshold=args.threshold,
            max_review_passes=args.max_review_passes,
            retry_max_chars=args.retry_max_chars,
        )
        records.append(record)
        initial_risk = "" if record.initial_risk is None else f"{record.initial_risk:.3f}"
        final_risk = "" if record.final_risk is None else f"{record.final_risk:.3f}"
        print(
            f"[{index}/{len(segments)}] initial={initial_risk} "
            f"final={final_risk} "
            f"passes={record.pass_count} unresolved={record.unresolved}",
            flush=True,
        )

    write_reports(
        output_dir,
        input_path,
        original_text,
        records,
        threshold=args.threshold,
        segment_max_chars=args.segment_max_chars,
        max_review_passes=args.max_review_passes,
    )
    print(f"Output directory: {output_dir}")
    print(f"Merged text: {output_dir / 'merged_revised.txt'}")
    print(f"AWAS packet: {output_dir / 'awas_review_packet.md'}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
