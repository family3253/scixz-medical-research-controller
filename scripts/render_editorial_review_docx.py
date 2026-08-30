#!/usr/bin/env python3
"""Render a full editorial peer-review report while preserving strict fusion audit data."""

from __future__ import annotations

import argparse
import importlib.util
import json
import zipfile
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
LANGUAGES = {"zh": "中文", "en": "English"}


def _load_base():
    path = ROOT / "scripts" / "render_final_review_docx.py"
    spec = importlib.util.spec_from_file_location("scixz_compact_review_renderer", path)
    if spec is None or spec.loader is None:
        raise RuntimeError("compact review renderer is unavailable")
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


BASE = _load_base()


def load_json(path: Path) -> Dict[str, Any]:
    payload = json.loads(path.read_text(encoding="utf-8-sig"))
    if not isinstance(payload, dict):
        raise ValueError(f"Expected a JSON object: {path}")
    return payload


def bilingual(value: Any, language: str, path: str, errors: Optional[List[str]] = None) -> str:
    errors = errors if errors is not None else []
    if not isinstance(value, dict) or not isinstance(value.get(language), str) or not value[language].strip():
        errors.append(f"{path}.{language} is required")
        return ""
    return value[language].strip()


def validate_profile(profile: Dict[str, Any], review: Dict[str, Any]) -> List[str]:
    errors: List[str] = []
    required = (
        "review_metadata", "editorial_decision_note", "reviewer_panel", "major_issue_profiles",
        "minor_issue_profiles", "recommended_revisions", "adversarial_stress_test", "author_questions",
        "dimension_scores", "reporting_completeness", "revision_roadmap", "recommended_revision_period",
    )
    for key in required:
        if key not in profile:
            errors.append(f"{key} is required")
    if errors:
        return errors
    metadata = profile["review_metadata"]
    if not isinstance(metadata, dict):
        errors.append("review_metadata must be an object")
    else:
        for key in ("review_type", "review_date", "main_manuscript"):
            value = metadata.get(key)
            if key == "review_date":
                if not isinstance(value, str) or not value.strip():
                    errors.append("review_metadata.review_date is required")
            else:
                for language in LANGUAGES:
                    bilingual(value, language, f"review_metadata.{key}", errors)
        if not isinstance(metadata.get("companion_files"), list):
            errors.append("review_metadata.companion_files must be a list")
    for language in LANGUAGES:
        bilingual(profile["editorial_decision_note"], language, "editorial_decision_note", errors)
        bilingual(profile["adversarial_stress_test"], language, "adversarial_stress_test", errors)
        bilingual(profile["reporting_completeness"], language, "reporting_completeness", errors)
        bilingual(profile["recommended_revision_period"], language, "recommended_revision_period", errors)
    major_ids = {str(item.get("id")) for item in review.get("major_concerns", []) if isinstance(item, dict)}
    minor_ids = {str(item.get("id")) for item in review.get("minor_concerns", []) if isinstance(item, dict)}
    for key, expected_ids in (("major_issue_profiles", major_ids), ("minor_issue_profiles", minor_ids)):
        items = profile.get(key)
        if not isinstance(items, list):
            errors.append(f"{key} must be a list")
            continue
        observed_ids = [str(item.get("concern_id")) for item in items if isinstance(item, dict)]
        if set(observed_ids) != expected_ids or len(observed_ids) != len(expected_ids):
            errors.append(f"{key} must map every corresponding concern id exactly once")
        for index, item in enumerate(items, 1):
            if not isinstance(item, dict):
                errors.append(f"{key}[{index}] must be an object")
                continue
            for field in (("title", "source", "acceptance_criteria") if key == "major_issue_profiles" else ("title",)):
                for language in LANGUAGES:
                    bilingual(item.get(field), language, f"{key}[{index}].{field}", errors)
    for key in ("reviewer_panel", "recommended_revisions", "author_questions", "dimension_scores", "revision_roadmap"):
        if not isinstance(profile.get(key), list) or not profile[key]:
            errors.append(f"{key} must be a non-empty list")
    return errors


def _set_repeat_table_header(row: Any) -> None:
    properties = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    properties.append(header)


def _table(document: Document, headers: List[str], rows: List[List[str]]) -> Any:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_repeat_table_header(table.rows[0])
    for cell, header in zip(table.rows[0].cells, headers):
        BASE._set_cell_text(cell, header, bold=True)
        BASE._shade(cell, "D9EAF7")
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            BASE._set_cell_text(cell, str(value))
    return table


def _label_paragraph(document: Document, label: str, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(label)
    run.bold = True
    paragraph.add_run(text)


def _numbered(document: Document, text: str) -> None:
    document.add_paragraph(text, style="List Number")


def _concern_map(review: Dict[str, Any], key: str) -> Dict[str, Dict[str, Any]]:
    return {str(item["id"]): item for item in review[key]}


def render(review: Dict[str, Any], profile: Dict[str, Any], language: str, output: Path) -> None:
    errors = BASE.validate(review) + validate_profile(profile, review)
    if errors:
        raise ValueError("; ".join(errors))
    document = Document()
    BASE._configure(document)
    title = "同行评审报告（中文版）" if language == "zh" else "Peer-Review Report (English)"
    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor(31, 78, 121)

    metadata = profile["review_metadata"]
    BASE._heading(document, "审阅稿件" if language == "zh" else "Materials Reviewed")
    labels = {
        "main": "主文稿：" if language == "zh" else "Main manuscript: ",
        "companions": "表格与图片文件：" if language == "zh" else "Companion tables/figures: ",
        "type": "审稿类型：" if language == "zh" else "Review type: ",
        "date": "审稿日期：" if language == "zh" else "Review date: ",
    }
    _label_paragraph(document, labels["main"], bilingual(metadata["main_manuscript"], language, "main_manuscript"))
    companions = [bilingual(item, language, "companion_files") for item in metadata["companion_files"]]
    _label_paragraph(document, labels["companions"], "; ".join(companions) if companions else ("无" if language == "zh" else "None"))
    _label_paragraph(document, labels["type"], bilingual(metadata["review_type"], language, "review_type"))
    _label_paragraph(document, labels["date"], metadata["review_date"])

    BASE._heading(document, "编辑决定" if language == "zh" else "Editorial Decision")
    decision_table = document.add_table(rows=1, cols=1)
    decision_table.style = "Table Grid"
    decision_cell = decision_table.cell(0, 0)
    BASE._shade(decision_cell, "FFF2CC")
    decision_text = bilingual(review["decision"], language, "decision")
    BASE._set_cell_text(decision_cell, decision_text, bold=True)
    document.add_paragraph(bilingual(profile["editorial_decision_note"], language, "editorial_decision_note"))

    BASE._heading(document, "审稿人配置" if language == "zh" else "Reviewer Panel")
    panel_headers = ["审稿人", "审稿视角", "建议", "信心评分"] if language == "zh" else ["Reviewer", "Perspective", "Recommendation", "Confidence"]
    panel_rows = []
    for item in profile["reviewer_panel"]:
        panel_rows.append([
            item["id"], bilingual(item["perspective"], language, "perspective"),
            bilingual(item["recommendation"], language, "recommendation"), str(item["confidence"]),
        ])
    _table(document, panel_headers, panel_rows)

    BASE._heading(document, "总体评价" if language == "zh" else "Overall Assessment")
    document.add_paragraph(bilingual(review["overall_assessment"], language, "overall_assessment"))
    BASE._heading(document, "主要优点" if language == "zh" else "Major Strengths")
    for item in review["strengths"]:
        _numbered(document, bilingual(item, language, "strength"))

    major_map = _concern_map(review, "major_concerns")
    BASE._heading(document, "必须修改的问题（P0：重新送审前必须解决）" if language == "zh" else "Mandatory Revisions (P0: required before re-review)")
    for item in profile["major_issue_profiles"]:
        concern = major_map[item["concern_id"]]
        heading = f"{concern['id']}. {bilingual(item['title'], language, 'title')}"
        BASE._heading(document, heading, level=2)
        _label_paragraph(document, "意见来源：" if language == "zh" else "Sources: ", bilingual(item["source"], language, "source"))
        document.add_paragraph(bilingual(concern["concern"], language, "concern"))
        _label_paragraph(document, "必须采取的措施：" if language == "zh" else "Required action: ", bilingual(concern["recommendation"], language, "recommendation"))
        _label_paragraph(document, "验收标准：" if language == "zh" else "Acceptance criterion: ", bilingual(item["acceptance_criteria"], language, "acceptance_criteria"))

    minor_map = _concern_map(review, "minor_concerns")
    BASE._heading(document, "必须修改的问题（P1：重要的次级问题）" if language == "zh" else "Important Secondary Revisions (P1)")
    for index, item in enumerate(profile["minor_issue_profiles"], 1):
        concern = minor_map[item["concern_id"]]
        paragraph = document.add_paragraph(style="List Number")
        title_run = paragraph.add_run(bilingual(item["title"], language, "title") + "：")
        title_run.bold = True
        paragraph.add_run(bilingual(concern["concern"], language, "concern") + " " + bilingual(concern["recommendation"], language, "recommendation"))

    BASE._heading(document, "建议修改的问题（P2/P3）" if language == "zh" else "Recommended Revisions (P2/P3)")
    for item in profile["recommended_revisions"]:
        BASE._bullet(document, bilingual(item, language, "recommended_revisions"))

    BASE._heading(document, "反方论证压力测试" if language == "zh" else "Adversarial Stress Test")
    document.add_paragraph(bilingual(profile["adversarial_stress_test"], language, "adversarial_stress_test"))
    BASE._heading(document, "作者必须回答的问题" if language == "zh" else "Questions the Authors Must Answer")
    for item in profile["author_questions"]:
        _numbered(document, bilingual(item, language, "author_question"))

    BASE._heading(document, "维度评分" if language == "zh" else "Dimension Scores")
    score_headers = ["维度", "得分", "评价"] if language == "zh" else ["Dimension", "Score", "Assessment"]
    score_rows = [[bilingual(item["dimension"], language, "dimension"), str(item["score"]), bilingual(item["assessment"], language, "assessment")] for item in profile["dimension_scores"]]
    _table(document, score_headers, score_rows)
    document.add_paragraph(bilingual(profile["reporting_completeness"], language, "reporting_completeness"))

    BASE._heading(document, "修订路线图" if language == "zh" else "Revision Roadmap")
    for phase in profile["revision_roadmap"]:
        BASE._heading(document, bilingual(phase["priority"], language, "priority"), level=2)
        for item in phase["items"]:
            document.add_paragraph("□ " + bilingual(item, language, "roadmap_item"))
    _label_paragraph(document, "建议修订周期：" if language == "zh" else "Recommended revision period: ", bilingual(profile["recommended_revision_period"], language, "recommended_revision_period"))

    BASE._heading(document, "附录A：外部 AI 审稿意见的处理" if language == "zh" else "Appendix A: External AI-Review Dispositions")
    for item in review["external_signal_integration"]:
        label = " / ".join(str(value) for value in (item.get("source_issue_id"), item["disposition"], item.get("manuscript_location")) if value)
        BASE._bullet(document, f"[{label}] " + bilingual(item["external_issue"], language, "external_issue"))
        document.add_paragraph(bilingual(item["rationale"], language, "rationale"))

    trace = review["synthesis_trace"]
    BASE._heading(document, "附录B：双路径一致性与分歧" if language == "zh" else "Appendix B: Cross-Branch Agreement and Disagreement")
    matrix_headers = ["ID", "分类", "来源问题", "核验结论"] if language == "zh" else ["ID", "Class", "Source issues", "Resolution"]
    matrix_rows = []
    for row in trace["agreement_disagreement_matrix"]:
        sources = ", ".join(row.get("local_issue_ids", []) + row.get("external_issue_ids", [])) or "-"
        matrix_rows.append([row["id"], row["classification"], sources, bilingual(row["resolution"], language, "resolution")])
    _table(document, matrix_headers, matrix_rows)
    scope = trace.get("evidence_scope")
    if isinstance(scope, dict):
        BASE._heading(document, "附录C：证据范围边界" if language == "zh" else "Appendix C: Evidence-Scope Boundary")
        document.add_paragraph(bilingual(scope["external_branch_limitations"], language, "external_branch_limitations"))
    BASE._heading(document, "限制与核验边界" if language == "zh" else "Limitations and Verification")
    for item in review["limitations"]:
        BASE._bullet(document, bilingual(item, language, "limitation"))

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    with zipfile.ZipFile(output) as archive:
        if "word/document.xml" not in archive.namelist():
            raise RuntimeError("DOCX validation failed: missing word/document.xml")
    Document(output)


def main(argv: Optional[List[str]] = None) -> int:
    parser = argparse.ArgumentParser(description="Render the full editorial peer-review structure into bilingual DOCX files.")
    parser.add_argument("--input", required=True)
    parser.add_argument("--structure-profile", required=True)
    parser.add_argument("--fusion-bundle")
    parser.add_argument("--output-dir", required=True)
    parser.add_argument("--stem", default="scixz_editorial_review")
    args = parser.parse_args(argv)
    review = load_json(Path(args.input))
    profile = load_json(Path(args.structure_profile))
    validation = BASE.validate(review)
    if args.fusion_bundle:
        fusion = load_json(Path(args.fusion_bundle))
        validation = BASE.validate(
            review,
            fusion.get("barrier", {}).get("external_issue_ids_requiring_disposition"),
            fusion.get("manuscript", {}).get("fingerprint"),
            fusion.get("barrier", {}).get("local_issue_ids_requiring_matrix_mapping"),
            fusion.get("evidence_manifest"),
        )
    validation += validate_profile(profile, review)
    if validation:
        parser.exit(2, "Full editorial review rendering blocked: " + "; ".join(validation) + "\n")
    destination = Path(args.output_dir)
    outputs = {language: destination / f"{args.stem}_{language}.docx" for language in LANGUAGES}
    for language, path in outputs.items():
        render(review, profile, language, path)
    print(json.dumps({"status": "COMPLETE", "outputs": {key: str(value) for key, value in outputs.items()}}, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
