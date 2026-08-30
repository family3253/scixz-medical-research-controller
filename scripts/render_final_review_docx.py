#!/usr/bin/env python3
"""Render a verified bilingual SciXZ final review JSON into Chinese and English DOCX files."""

from __future__ import annotations

import argparse
import json
import zipfile
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

try:
    from scripts.private_artifact_guard import ensure_private_output_path
except ImportError:
    from private_artifact_guard import ensure_private_output_path


LANGUAGES = {"zh": "中文", "en": "English"}
REQUIRED = ("metadata", "decision", "overall_assessment", "strengths", "major_concerns", "minor_concerns", "external_signal_integration", "limitations")


def _text(value: Any, language: str, path: str, errors: List[str]) -> str:
    if not isinstance(value, dict) or not isinstance(value.get(language), str) or not value[language].strip():
        errors.append(f"{path}.{language} is required")
        return ""
    return value[language].strip()


def validate(review: Dict[str, Any], expected_external_issue_ids: Optional[List[str]] = None, expected_fingerprint: Optional[str] = None, expected_local_issue_ids: Optional[List[str]] = None, expected_evidence_manifest: Optional[Dict[str, Any]] = None) -> List[str]:
    errors: List[str] = []
    for key in REQUIRED:
        if key not in review or review.get(key) is None:
            errors.append(f"{key} is required")
    if errors:
        return errors
    metadata = review.get("metadata")
    if not isinstance(metadata, dict):
        errors.append("metadata must be an object")
    else:
        for key in ("manuscript_title", "review_scope"):
            if not isinstance(metadata.get(key), str) or not metadata[key].strip():
                errors.append(f"metadata.{key} is required")
        if expected_fingerprint and metadata.get("manuscript_fingerprint") != expected_fingerprint:
            errors.append("metadata.manuscript_fingerprint must match the fusion bundle")
    if not isinstance(review.get("external_signal_integration"), list):
        errors.append("external_signal_integration must be a list")
        return errors
    for language in LANGUAGES:
        _text(review["decision"], language, "decision", errors)
        _text(review["overall_assessment"], language, "overall_assessment", errors)
        for collection in ("strengths", "major_concerns", "minor_concerns", "limitations"):
            if not isinstance(review.get(collection), list):
                errors.append(f"{collection} must be a list")
                continue
            for index, item in enumerate(review[collection], 1):
                if collection in {"major_concerns", "minor_concerns"}:
                    if not isinstance(item, dict) or not str(item.get("id", "")).strip() or not str(item.get("location", "")).strip():
                        errors.append(f"{collection}[{index}] needs id and location")
                    else:
                        _text(item.get("concern"), language, f"{collection}[{index}].concern", errors)
                        _text(item.get("recommendation"), language, f"{collection}[{index}].recommendation", errors)
                else:
                    _text(item, language, f"{collection}[{index}]", errors)
        for index, item in enumerate(review.get("external_signal_integration", []), 1):
            if not isinstance(item, dict) or item.get("disposition") not in {"incorporated", "incorporated-with-revision", "rejected", "unresolved"}:
                errors.append(f"external_signal_integration[{index}] needs an allowed disposition")
            else:
                _text(item.get("external_issue"), language, f"external_signal_integration[{index}].external_issue", errors)
                _text(item.get("rationale"), language, f"external_signal_integration[{index}].rationale", errors)
    if expected_external_issue_ids is not None:
        integrations = review.get("external_signal_integration", [])
        observed: List[str] = []
        for index, item in enumerate(integrations, 1):
            if not isinstance(item, dict):
                continue
            source_id = item.get("source_issue_id")
            if not isinstance(source_id, str) or not source_id.strip():
                errors.append(f"external_signal_integration[{index}].source_issue_id is required for parallel fusion")
            else:
                observed.append(source_id.strip())
            if not isinstance(item.get("manuscript_location"), str) or not item["manuscript_location"].strip():
                errors.append(f"external_signal_integration[{index}].manuscript_location is required for parallel fusion")
        if len(observed) != len(set(observed)):
            errors.append("external issue dispositions contain duplicate source_issue_id values")
        missing = sorted(set(expected_external_issue_ids) - set(observed))
        extra = sorted(set(observed) - set(expected_external_issue_ids))
        if missing:
            errors.append("external issue dispositions are missing: " + ", ".join(missing))
        if extra:
            errors.append("external issue dispositions contain unknown ids: " + ", ".join(extra))
        trace = review.get("synthesis_trace")
        if not isinstance(trace, dict):
            errors.append("synthesis_trace is required for parallel fusion")
        else:
            declared = trace.get("external_issue_ids_expected")
            if not isinstance(declared, list) or set(declared) != set(expected_external_issue_ids) or len(declared) != len(expected_external_issue_ids):
                errors.append("synthesis_trace.external_issue_ids_expected must match the fusion bundle")
            matrix = trace.get("agreement_disagreement_matrix")
            if not isinstance(matrix, list) or not matrix:
                errors.append("synthesis_trace.agreement_disagreement_matrix must be a non-empty list")
            else:
                allowed = {"agreement", "complementary", "disagreement", "local-only", "external-only"}
                matrix_local_ids: List[str] = []
                matrix_external_ids: List[str] = []
                for index, row in enumerate(matrix, 1):
                    if not isinstance(row, dict) or not str(row.get("id", "")).strip():
                        errors.append(f"synthesis_trace.agreement_disagreement_matrix[{index}] needs id")
                        continue
                    if row.get("classification") not in allowed:
                        errors.append(f"synthesis_trace.agreement_disagreement_matrix[{index}] has invalid classification")
                    for key in ("local_issue_ids", "external_issue_ids"):
                        if not isinstance(row.get(key), list):
                            errors.append(f"synthesis_trace.agreement_disagreement_matrix[{index}].{key} must be a list")
                    if isinstance(row.get("local_issue_ids"), list):
                        matrix_local_ids.extend(str(value).strip() for value in row["local_issue_ids"] if str(value).strip())
                    if isinstance(row.get("external_issue_ids"), list):
                        matrix_external_ids.extend(str(value).strip() for value in row["external_issue_ids"] if str(value).strip())
                    for language in LANGUAGES:
                        _text(row.get("resolution"), language, f"synthesis_trace.agreement_disagreement_matrix[{index}].resolution", errors)
                for label, observed_ids, expected_ids in (
                    ("external", matrix_external_ids, expected_external_issue_ids),
                    ("local", matrix_local_ids, expected_local_issue_ids),
                ):
                    if expected_ids is None:
                        continue
                    if len(observed_ids) != len(set(observed_ids)):
                        errors.append(f"agreement/disagreement matrix contains duplicate {label} issue ids")
                    missing_ids = sorted(set(expected_ids) - set(observed_ids))
                    unknown_ids = sorted(set(observed_ids) - set(expected_ids))
                    if missing_ids:
                        errors.append(f"agreement/disagreement matrix is missing {label} issue ids: " + ", ".join(missing_ids))
                    if unknown_ids:
                        errors.append(f"agreement/disagreement matrix contains unknown {label} issue ids: " + ", ".join(unknown_ids))
            if not isinstance(trace.get("dissenting_sources"), list):
                errors.append("synthesis_trace.dissenting_sources must be a list")
            companions = (expected_evidence_manifest or {}).get("companion_evidence", [])
            if companions:
                scope = trace.get("evidence_scope")
                if not isinstance(scope, dict):
                    errors.append("synthesis_trace.evidence_scope is required when companion evidence exists")
                else:
                    expected_companion_fingerprints = [item.get("fingerprint") for item in companions]
                    if scope.get("shared_uploaded_pdf_fingerprint") != expected_fingerprint:
                        errors.append("synthesis_trace.evidence_scope shared PDF fingerprint must match the fusion bundle")
                    if scope.get("branch_scopes_identical") is not False:
                        errors.append("synthesis_trace.evidence_scope must disclose that branch scopes are not identical")
                    observed_companions = scope.get("companion_evidence_fingerprints")
                    if not isinstance(observed_companions, list) or set(observed_companions) != set(expected_companion_fingerprints) or len(observed_companions) != len(expected_companion_fingerprints):
                        errors.append("synthesis_trace.evidence_scope companion fingerprints must match the fusion bundle")
                    for language in LANGUAGES:
                        _text(scope.get("external_branch_limitations"), language, "synthesis_trace.evidence_scope.external_branch_limitations", errors)
    return errors


def _shade(cell: Any, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    for child in list(properties):
        if child.tag == qn("w:shd"):
            properties.remove(child)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), fill)
    elements_after_shading = {qn(name) for name in ("w:noWrap", "w:tcMar", "w:textDirection", "w:tcFitText", "w:vAlign", "w:hideMark")}
    for index, child in enumerate(properties):
        if child.tag in elements_after_shading:
            properties.insert(index, shading)
            break
    else:
        properties.append(shading)


def _set_cell_text(cell: Any, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def _configure(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    zoom = document.settings.element.find(qn("w:zoom"))
    if zoom is not None:
        zoom.set(qn("w:percent"), "100")


def _heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_heading(text, level=level)
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(5)


def _bullet(document: Document, text: str) -> None:
    document.add_paragraph(text, style="List Bullet")


def render(review: Dict[str, Any], language: str, output: Path) -> None:
    errors = validate(review)
    if errors:
        raise ValueError("; ".join(errors))
    label = LANGUAGES[language]
    document = Document()
    _configure(document)
    title = "SciXZ 最终同行评审意见" if language == "zh" else "SciXZ Final Peer-Review Report"
    subtitle = "已核验的外部 AI 审稿信号仅作辅助参考" if language == "zh" else "Verified external AI-review signals are advisory evidence only"
    title_p = document.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor(31, 78, 121)
    sub_p = document.add_paragraph(subtitle)
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.runs[0].italic = True

    meta = review["metadata"]
    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    labels = [("Manuscript" if language == "en" else "稿件", meta.get("manuscript_title", "")), ("Decision" if language == "en" else "总体建议", _text(review["decision"], language, "decision", [])), ("Review scope" if language == "en" else "审稿范围", meta.get("review_scope", "")), ("Language" if language == "en" else "报告语言", label)]
    for left, right in labels:
        cells = table.add_row().cells
        _set_cell_text(cells[0], str(left), bold=True)
        _set_cell_text(cells[1], str(right))
        _shade(cells[0], "D9EAF7")

    _heading(document, "Overall Assessment" if language == "en" else "总体评价")
    document.add_paragraph(_text(review["overall_assessment"], language, "overall_assessment", []))
    _heading(document, "Strengths" if language == "en" else "主要优点")
    for item in review["strengths"]:
        _bullet(document, _text(item, language, "strengths", []))

    for key, heading in (("major_concerns", "Major Concerns" if language == "en" else "主要问题"), ("minor_concerns", "Minor Concerns" if language == "en" else "次要问题")):
        _heading(document, heading)
        concerns = review[key]
        if not concerns:
            document.add_paragraph("None." if language == "en" else "无。")
            continue
        table = document.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ["ID", "Location" if language == "en" else "位置", "Concern" if language == "en" else "问题", "Recommendation" if language == "en" else "建议"]
        for cell, header in zip(table.rows[0].cells, headers):
            _set_cell_text(cell, header, bold=True)
            _shade(cell, "D9EAF7")
        for item in concerns:
            cells = table.add_row().cells
            values = [item["id"], item["location"], _text(item["concern"], language, "concern", []), _text(item["recommendation"], language, "recommendation", [])]
            for cell, value in zip(cells, values):
                _set_cell_text(cell, str(value))

    _heading(document, "External AI-Review Integration" if language == "en" else "外部 AI 审稿意见的处理")
    if not review["external_signal_integration"]:
        document.add_paragraph("No external AI-review signal was used." if language == "en" else "未使用外部 AI 审稿信号。")
    for item in review["external_signal_integration"]:
        issue = _text(item["external_issue"], language, "external_issue", [])
        rationale = _text(item["rationale"], language, "rationale", [])
        prefix = item.get("source_issue_id")
        location = item.get("manuscript_location")
        label = " / ".join(str(value) for value in (prefix, item["disposition"], location) if value)
        document.add_paragraph(f"[{label}] {issue}", style="List Bullet")
        document.add_paragraph(rationale)
    trace = review.get("synthesis_trace")
    if isinstance(trace, dict) and trace.get("agreement_disagreement_matrix"):
        _heading(document, "Cross-Branch Agreement and Disagreement" if language == "en" else "双路径一致性与分歧")
        matrix = document.add_table(rows=1, cols=4)
        matrix.style = "Table Grid"
        matrix.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ["ID", "Class" if language == "en" else "分类", "Source issues" if language == "en" else "来源问题", "Resolution" if language == "en" else "核验结论"]
        for cell, header in zip(matrix.rows[0].cells, headers):
            _set_cell_text(cell, header, bold=True)
            _shade(cell, "D9EAF7")
        for row in trace["agreement_disagreement_matrix"]:
            cells = matrix.add_row().cells
            sources = ", ".join(row.get("local_issue_ids", []) + row.get("external_issue_ids", [])) or "-"
            values = [row.get("id", ""), row.get("classification", ""), sources, _text(row.get("resolution"), language, "resolution", [])]
            for cell, value in zip(cells, values):
                _set_cell_text(cell, str(value))
        scope = trace.get("evidence_scope")
        if isinstance(scope, dict):
            _heading(document, "Evidence-Scope Boundary" if language == "en" else "证据范围边界")
            document.add_paragraph(_text(scope.get("external_branch_limitations"), language, "external_branch_limitations", []))
    _heading(document, "Limitations and Verification" if language == "en" else "限制与核验边界")
    for item in review["limitations"]:
        _bullet(document, _text(item, language, "limitations", []))

    output = ensure_private_output_path(output)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    with zipfile.ZipFile(output) as archive:
        if "word/document.xml" not in archive.namelist():
            raise RuntimeError("DOCX validation failed: missing word/document.xml")
    Document(output)


def main(argv: Optional[List[str]] = None) -> int:
    parser = argparse.ArgumentParser(description="Render Chinese and English final-review DOCX files from a verified bilingual JSON review.")
    parser.add_argument("--input", required=True)
    parser.add_argument("--output-dir", required=True)
    parser.add_argument("--stem", default="scixz_final_review")
    parser.add_argument("--fusion-bundle", help="Enable strict parallel-fusion coverage checks against this bundle")
    args = parser.parse_args(argv)
    review = json.loads(Path(args.input).read_text(encoding="utf-8-sig"))
    if not isinstance(review, dict):
        parser.error("input must be a JSON object")
    expected_ids: Optional[List[str]] = None
    expected_local_ids: Optional[List[str]] = None
    expected_fingerprint: Optional[str] = None
    expected_evidence_manifest: Optional[Dict[str, Any]] = None
    if args.fusion_bundle:
        fusion = json.loads(Path(args.fusion_bundle).read_text(encoding="utf-8-sig"))
        if not isinstance(fusion, dict):
            parser.error("fusion bundle must be a JSON object")
        expected_ids = fusion.get("barrier", {}).get("external_issue_ids_requiring_disposition")
        expected_local_ids = fusion.get("barrier", {}).get("local_issue_ids_requiring_matrix_mapping")
        expected_fingerprint = fusion.get("manuscript", {}).get("fingerprint")
        expected_evidence_manifest = fusion.get("evidence_manifest")
        if not isinstance(expected_ids, list) or not expected_ids:
            parser.error("fusion bundle lacks external issue ids")
        if not isinstance(expected_local_ids, list) or not expected_local_ids:
            parser.error("fusion bundle lacks local issue ids")
        if not isinstance(expected_fingerprint, str) or not expected_fingerprint.startswith("sha256:"):
            parser.error("fusion bundle lacks a manuscript fingerprint")
        if not isinstance(expected_evidence_manifest, dict):
            parser.error("fusion bundle lacks an evidence manifest")
    errors = validate(review, expected_ids, expected_fingerprint, expected_local_ids, expected_evidence_manifest)
    if errors:
        parser.exit(2, "Final review DOCX rendering blocked: " + "; ".join(errors) + "\n")
    destination = Path(args.output_dir)
    outputs = {language: destination / f"{args.stem}_{language}.docx" for language in LANGUAGES}
    for language, path in outputs.items():
        render(review, language, path)
    print(json.dumps({"status": "COMPLETE", "outputs": {key: str(value) for key, value in outputs.items()}}, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
